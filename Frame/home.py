from tkinter import *
from tkinter import scrolledtext
from tkinter import font
from PIL import ImageTk,Image
from PyPDF2 import pdf
from Frame import login
from tkinter import filedialog,scrolledtext,ttk
from tkinter.messagebox import askokcancel, showinfo, WARNING
from PIL import ImageTk,Image
from tkinter.scrolledtext import ScrolledText
import PyPDF2,os,docx,pyperclip
import requests
class Mainhome:
    def __init__(self):
        w=Tk()
        w.geometry('1400x700') 
        w.configure(bg='#000000')
        w.resizable(False, False) 
        w.title("Filo")
        w.iconbitmap(r'Frame/home_img/icon.ico')
        def default_home():
            homeframe=Frame(w,width=1400,height=658,bg='#262626')
            glabel=Label(w,text='Home',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42) 
            homeframe.place(x=0,y=42)

        def toggle_win():
            global menu
            menu=Frame(w,width=300,height=700,bg='#D2E6FB')
            menu.place(x=0,y=0)

            def bttn(x,y,text,bcolor,fcolor,cmd):
                def on_entera(e):
                    myButton1['background']=bcolor
                    myButton1['foreground']='#262626'
                
                def on_leavea(e):
                    myButton1['background']=fcolor
                    myButton1['foreground']='#262626'

                myButton1= Button(menu,text=text,width=25,height=1,fg='#262626',border=0,font=("Poppins",15),bg=fcolor,activebackground='#262626',activeforeground=bcolor,command=cmd)
                myButton1.bind("<Enter>",on_entera)
                myButton1.bind("<Enter>",on_leavea)
                myButton1.place(x=x,y=y)

            def pdf_code():
                #function to merge pdfs
                def pdfmerge():
                    mergeframe=Frame(pdfframe,width=1400,height=658,bg='#000000')
                    mergeframe.place(x=700,y=0)
                    Label(mergeframe,text='PDF Merge',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=250,y=0)
                    global pdfmergecount
                    pdfmergecount=0
                    def selectfile():
                        global pdfmergecount,mergefile1,mergefile2
                        if(pdfmergecount!=2):
                            if pdfmergecount==0:
                                mergefile1=filedialog.askopenfilename(initialdir="C:\\",title="Open",filetypes=(("PDF Files","*.pdf"),("All files","*.*")))
                                if mergefile1!='' and mergefile1:
                                    pdfmergecount+=1
                                    t1=Entry(mergeframe,bd=0,highlightthickness=0,font=("Poppins",16))
                                    t1.place(x=100,y=230,width=650)
                                    t1.insert(0,os.path.basename(mergefile1))
                                    t1.configure(state='readonly',readonlybackground="#000000",foreground="#FFFFFF")
                            else:
                                mergefile2=filedialog.askopenfilename(initialdir="C:\\",title="Open",filetypes=(("PDF Files","*.pdf"),("All files","*.*")))
                                if mergefile2!='':
                                    pdfmergecount+=1
                                    t1=Entry(mergeframe,bd=0,highlightthickness=0,font=("Poppins",16))
                                    t1.place(x=100,y=330,width=650)
                                    t1.insert(0,os.path.basename(mergefile2))
                                    t1.configure(state='readonly',readonlybackground="#000000",foreground="#FFFFFF")
                        else:
                            showinfo(title='Error',message='Maximum two files Supported.',icon=WARNING)
                    def finalmerge():
                        if pdfmergecount!=2:
                            showinfo(title='Error',message='Merge requires two files.',icon=WARNING)
                        else:
                            pdf1File = open(mergefile1, 'rb')
                            pdf2File = open(mergefile2, 'rb')
                            # File Readers
                            pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
                            pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
                            #File writer
                            if pdf1Reader.isEncrypted==False and pdf2Reader.isEncrypted==False:
                                pdfWriter = PyPDF2.PdfFileWriter()
                                for pageNum in range(pdf1Reader.numPages):
                                    pageObj = pdf1Reader.getPage(pageNum)
                                    pdfWriter.addPage(pageObj)
                                
                                for pageNum in range(pdf2Reader.numPages):
                                    pageObj = pdf2Reader.getPage(pageNum)
                                    pdfWriter.addPage(pageObj)
                                pdfPath = filedialog.asksaveasfilename(defaultextension = "*.pdf", filetypes = (("PDF Files", "*.pdf"),))
                                if pdfPath: #If the user only selects the save file location
                                    pdfOutputFile = open(pdfPath, 'wb')
                                    pdfWriter.write(pdfOutputFile)
                                    pdfOutputFile.close()
                                    pdf1File.close()
                                    pdf2File.close()
                                    showinfo(title='Success', message='File Merged Successfully!', icon='info')
                            else:
                                showinfo(title='Error', message='Encrypted Files Not Supported.', icon=WARNING)
                            pdfmerge()                           

                    mergebtn = Button(mergeframe,text="Merge & Save",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = finalmerge,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=450,y=558,width=200)
                    selectfilebtn = Button(mergeframe,text="Select File",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = selectfile,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=100,y=100,width=500)

                #funtion to split pdfs
                def pdfsplit():
                    print("split pdfs here")

                menu.destroy()
                glabel=Label(w,text='PDF Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                pdfframe=Frame(w,width=1400,height=658,bg='#262626')
                pdfframe.place(x=0,y=42)
                #merge pdf button
                merge_img = PhotoImage(file = f"Frame/home_img/merge.png")
                label = Label(image=merge_img)
                label.image=merge_img
                mergebtn = Button(pdfframe,image = merge_img,borderwidth = 0,highlightthickness = 0,command = pdfmerge,background="#262626",activebackground="#262626",relief = "flat").place(x=130,y=240,width=129,height=119)
                
                #split pdf button
                split_img = PhotoImage(file = f"Frame/home_img/split.png")
                label = Label(image=split_img)
                label.image=split_img
                splitbtn = Button(pdfframe,image = split_img,borderwidth = 0,highlightthickness = 0,command = pdfsplit,background="#262626",activebackground="#262626",relief = "flat").place(x=330,y=240,width=129,height=119)
                pdfmerge()

            def word_code():
                menu.destroy()
                def textodoc():
                    doctext=textentry.get("1.0",END)
                    import docx,pyperclip
                    doc=docx.Document()
                    header=headerentry.get()
                    if header!='':
                        doc.add_header(header,0)
                    text=doctext.split('\n')
                    for i in range(len(text)):
                        doc.add_paragraph(text[i])
                        doc.paragraphs[i].runs[1].style='QuoteChar'
                    wordPath = filedialog.asksaveasfilename(defaultextension = "*.docx", filetypes = (("Word Files", "*.docx"),))
                    print(wordPath)
                    if wordPath:
                        doc.save(wordPath)
                    showinfo(title='Success', message='File created Successfully!', icon='info')


                    
                wordframe=Frame(w,width=1400,height=658,bg='#262626')
                glabel=Label(w,text='Word Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                wordframe.place(x=0,y=42)
                #entry of header
                Label(wordframe,text='Header',font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=600,y=0)
                headerentry=Entry(wordframe,font=("Poppins",20),background="#FFFFFF")
                headerentry.place(x=600,y=50,width=790,height=60)
                #content entry
                textentry=ScrolledText(wordframe,font=("Poppins",20),background="#FFFFFF")
                textentry.place(x=600,y=120,width=790,height=508)

                # Label(wordframe,text='Font Size',font=("Poppins",18),background='#262626',foreground='#FFFFFF').place(x=335,y=160)
                # fontsize=Entry(wordframe,font=("Poppins",18),background="#FFFFFF")
                # fontsize.place(x=340,y=200,width=80,height=40)
                
                # Label(wordframe,text='Font Style',font=("Poppins",18),background='#262626',foreground='#FFFFFF').place(x=100,y=160)
                # fontstyleval = StringVar()
                # fontstyle = ttk.Combobox(wordframe,textvariable = fontstyleval,foreground="#000000",font=("Poppins",10),state='readonly')
                # fontstyle.place(x=100,y=200,height=40)
                # fontstyle['values'] = ('Cambria','Comic Sans MS','Calibri','Cavolini','Arial')
                # fontstyle.current(1)

                doc_img = PhotoImage(file = f"Frame/home_img/doc.png")
                label = Label(image=doc_img)
                label.image=doc_img
                Button(wordframe,image=doc_img,borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = textodoc,background="#262626",foreground="#000000",activebackground="#262626",relief = "flat").place(x=230,y=300,width=129,height=119)

            def home():
                menu.destroy()
                default_home()

            def excel_code():
                menu.destroy()
                excelframe=Frame(w,width=1400,height=658,bg='#262626')
                excelframe.place(x=0,y=42)
                                
                def details():
                    api_key = "fc997859c662f09ea9dac60a3c71ecbc"
                    apiurl = "http://api.openweathermap.org/data/2.5/weather?"
                    city_name = input("Enter city name : ")

                    url = apiurl + "appid=" + api_key + "&q=" + city_name
                    response = requests.get(url)
                    x = response.json()

                    #checking if city is found or not
                    if x["cod"] != "404":

                        y = x["main"]
                        z = x["weather"]
                        a = x['coord']
                        b = x['wind']
                        c = x['sys'] 

                        latitude=a['lat']
                        longitude=a['lon']
                        temp = y["temp"]
                        pressure = y["pressure"]
                        humidity = y["humidity"]
                        country=c['country']
                        desc = z[0]["description"]
                        wind_speed=b['speed']
                    
                    else:
                        print(" City Not Found ")
                search_img = PhotoImage(file = f"Frame/home_img/search.png")
                label = Label(image=search_img)
                label.image=search_img
                Button(excelframe,borderwidth = 0,image=search_img,highlightthickness = 0,font=("Poppins",15),command = None,background="#262626",foreground="#000000",activebackground="#262626",relief = "flat").place(x=830,y=50,width=37,height=39)
                
                headerentry=Entry(excelframe,font=("Poppins",20),background="#FFFFFF")
                headerentry.place(x=600,y=50,width=790,height=60)

                Button(excelframe,borderwidth = 0,text="Submit",highlightthickness = 0,font=("Poppins",15),command = None,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=830,y=250)
                glabel=Label(w,text='Excel Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                details()


            def csv_code():
                menu.destroy()
                f6=Frame(w,width=1400,height=658,bg='#262626')
                glabel=Label(w,text='CSV Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                f6.place(x=0,y=42)

            def services_code():
                menu.destroy()
                f7=Frame(w,width=1400,height=658,bg='#262626')
                f7.place(x=0,y=42)
                glabel=Label(w,text='Automated Services',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)

            def logout():
                w.destroy()
                login.Main_window('login')

            bttn(0,80,'Home','#D2E6FB','#D2E6FB',home)
            bttn(0,150,'Manage PDF','#D2E6FB','#D2E6FB',pdf_code)
            bttn(0,220,'Manage Word','#D2E6FB','#D2E6FB',word_code)
            bttn(0,290,'Manage Excel','#D2E6FB','#D2E6FB',excel_code)
            bttn(0,360,'Manage CSV','#D2E6FB','#D2E6FB',csv_code)
            bttn(0,430,'Services','#D2E6FB','#D2E6FB',services_code)
            bttn(0,650,'Logout','#D2E6FB','#D2E6FB',logout)
                
            def dele():
                menu.destroy()

                
            global menubtn
            menubtn=ImageTk.PhotoImage(Image.open('Frame/home_img/close.png'))

            Button(menu,image=menubtn,command=dele,border=0,background='#D2E6FB',activebackground='#D2E6FB').place(x=5,y=10)

        default_home()

        img1=ImageTk.PhotoImage(Image.open('Frame/home_img/open.png'))
        Button(w,command=toggle_win,image=img1,border=0,bg='#000000',activebackground='#000000').place(x=5,y=5)

        app_width=1400 
        app_height=700
        screen_width = w.winfo_screenwidth()
        screen_height = w.winfo_screenheight()
        x=(screen_width / 2) - (app_width/2)
        y=(screen_height / 2) - (app_height/2)
        w.geometry(f'{app_width}x{app_height}+{int(x)}+{int(y)}')
        w.deiconify()
        w.mainloop()